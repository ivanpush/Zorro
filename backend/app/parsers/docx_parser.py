"""DOCX document parser using python-docx."""

import re
from datetime import datetime
from pathlib import Path
from uuid import uuid4

import structlog
from docx import Document
from docx.text.paragraph import Paragraph as DocxParagraph

from app.models.document import (
    DocObj,
    DocumentMetadata,
    Figure,
    Paragraph,
    Reference,
    Section,
    Sentence,
)
from .base import BaseParser

logger = structlog.get_logger()


class DocxParser(BaseParser):
    """Parser for DOCX documents."""

    def __init__(self):
        # Regex for sentence splitting (handles common abbreviations)
        self.sentence_splitter = re.compile(
            r'(?<=[.!?])\s+(?=[A-Z])|'  # Normal sentence end
            r'(?<=\w)\.\s*$'  # End of text
        )
        # Pattern for figure captions
        self.figure_pattern = re.compile(
            r'^(Figure|Fig\.?)\s+\d+[:\.]?\s*',
            re.IGNORECASE
        )
        # Pattern for references section
        self.references_pattern = re.compile(
            r'^(References?|Bibliography|Works?\s+Cited)',
            re.IGNORECASE
        )

    async def parse(self, file_path: Path, title_override: str | None = None) -> DocObj:
        """Parse a DOCX file into a DocObj."""
        logger.info("parsing_docx", file_path=str(file_path))

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"Failed to open DOCX file: {e}")

        # Extract components
        sections = self._extract_sections(doc)
        paragraphs, figures = self._extract_paragraphs_and_figures(doc, sections)
        references = self._extract_references(paragraphs)
        metadata = self._extract_metadata(doc, paragraphs)

        # Determine title
        title = title_override or self._extract_title(doc, sections)

        doc_obj = DocObj(
            document_id=str(uuid4()),
            filename=file_path.name,
            type="docx",
            title=title,
            sections=sections,
            paragraphs=paragraphs,
            figures=figures,
            references=references,
            metadata=metadata,
            created_at=datetime.utcnow()
        )

        logger.info(
            "docx_parsed",
            doc_id=doc_obj.document_id,
            sections=len(sections),
            paragraphs=len(paragraphs),
            figures=len(figures),
            references=len(references)
        )

        return doc_obj

    def _extract_sections(self, doc: Document) -> list[Section]:
        """Extract sections based on heading styles."""
        sections = []
        current_section_paragraphs = []
        section_index = 0

        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                # Save previous section if exists
                if current_section_paragraphs or sections:
                    if sections:
                        sections[-1].paragraph_ids = current_section_paragraphs
                    current_section_paragraphs = []

                # Extract heading level
                level = 1
                if 'Heading' in para.style.name:
                    try:
                        level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1

                sections.append(Section(
                    section_id=f"sec_{section_index:03d}",
                    section_index=section_index,
                    section_title=para.text.strip(),
                    level=level,
                    paragraph_ids=[]
                ))
                section_index += 1

        # Add remaining paragraphs to last section
        if sections and current_section_paragraphs:
            sections[-1].paragraph_ids = current_section_paragraphs

        # If no sections found, create a default one
        if not sections:
            sections.append(Section(
                section_id="sec_000",
                section_index=0,
                section_title=None,
                level=1,
                paragraph_ids=[]
            ))

        return sections

    def _extract_paragraphs_and_figures(
        self,
        doc: Document,
        sections: list[Section]
    ) -> tuple[list[Paragraph], list[Figure]]:
        """Extract paragraphs with sentences and detect figures."""
        paragraphs = []
        figures = []
        para_index = 0
        fig_index = 0
        current_section_idx = 0
        section_para_ids = []

        for docx_para in doc.paragraphs:
            # Skip empty paragraphs
            if not docx_para.text.strip():
                continue

            # Skip heading paragraphs
            if docx_para.style.name.startswith('Heading'):
                # Save paragraphs to previous section
                if current_section_idx < len(sections):
                    sections[current_section_idx].paragraph_ids = section_para_ids
                    section_para_ids = []
                    current_section_idx += 1
                continue

            para_id = f"p_{para_index:03d}"
            para_text = docx_para.text.strip()

            # Check if this is a figure caption
            if self.figure_pattern.match(para_text):
                figure = Figure(
                    figure_id=f"fig_{fig_index:03d}",
                    figure_index=fig_index,
                    caption=para_text,
                    caption_paragraph_id=para_id,
                    after_paragraph_id=paragraphs[-1].paragraph_id if paragraphs else None,
                    extraction_method="inline"
                )
                figures.append(figure)
                fig_index += 1

            # Extract sentences
            sentences = self._split_into_sentences(para_text, para_id)

            # Get XML path for export
            xml_path = self._get_xml_path(docx_para)

            # Determine section
            section_id = None
            if current_section_idx < len(sections):
                section_id = sections[current_section_idx].section_id
                section_para_ids.append(para_id)

            paragraph = Paragraph(
                paragraph_id=para_id,
                section_id=section_id,
                paragraph_index=para_index,
                text=para_text,
                sentences=sentences,
                xml_path=xml_path
            )
            paragraphs.append(paragraph)
            para_index += 1

        # Add remaining paragraphs to last section
        if current_section_idx < len(sections) and section_para_ids:
            sections[current_section_idx].paragraph_ids = section_para_ids

        # Also check for floating images and text box images
        figures.extend(self._extract_floating_figures(doc, paragraphs))

        return paragraphs, figures

    def _extract_floating_figures(
        self,
        doc: Document,
        paragraphs: list[Paragraph]
    ) -> list[Figure]:
        """Extract figures from floating shapes, text boxes, and tables."""
        figures = []
        fig_index = len([p for p in paragraphs if self.figure_pattern.match(p.text)])

        # Check for DrawingML shapes (floating images)
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                # This document contains images
                figure = Figure(
                    figure_id=f"fig_{fig_index:03d}",
                    figure_index=fig_index,
                    caption=None,
                    caption_paragraph_id=None,
                    after_paragraph_id=None,
                    extraction_method="float"
                )
                figures.append(figure)
                fig_index += 1

        # Check text boxes for images and captions
        for element in doc.element.body:
            if element.tag.endswith('}txbxContent'):
                text_content = element.text_content() if hasattr(element, 'text_content') else ""
                if self.figure_pattern.match(text_content):
                    figure = Figure(
                        figure_id=f"fig_{fig_index:03d}",
                        figure_index=fig_index,
                        caption=text_content,
                        caption_paragraph_id=None,
                        after_paragraph_id=None,
                        extraction_method="textbox"
                    )
                    figures.append(figure)
                    fig_index += 1

        return figures

    def _split_into_sentences(self, text: str, para_id: str) -> list[Sentence]:
        """Split paragraph text into sentences with character offsets."""
        sentences = []

        # Handle edge cases
        if not text:
            return sentences

        # Split by common sentence endings
        parts = re.split(r'([.!?]+)\s+', text)

        current_pos = 0
        sentence_index = 0
        current_sentence = ""

        for i, part in enumerate(parts):
            if not part:
                continue

            if re.match(r'^[.!?]+$', part):
                # This is punctuation
                current_sentence += part
                if i < len(parts) - 1:  # Not the last part
                    sentence_text = current_sentence.strip()
                    if sentence_text:
                        sentences.append(Sentence(
                            sentence_id=f"{para_id}_s_{sentence_index:03d}",
                            paragraph_id=para_id,
                            sentence_index=sentence_index,
                            text=sentence_text,
                            start_char=current_pos,
                            end_char=current_pos + len(sentence_text)
                        ))
                        current_pos = text.find(parts[i + 1] if i + 1 < len(parts) else "", current_pos + len(sentence_text))
                        sentence_index += 1
                        current_sentence = ""
            else:
                # This is sentence content
                if current_sentence and not current_sentence.endswith(' '):
                    current_sentence += ' '
                current_sentence += part

        # Add any remaining text as a sentence
        if current_sentence.strip():
            sentence_text = current_sentence.strip()
            sentences.append(Sentence(
                sentence_id=f"{para_id}_s_{sentence_index:03d}",
                paragraph_id=para_id,
                sentence_index=sentence_index,
                text=sentence_text,
                start_char=current_pos,
                end_char=len(text)
            ))

        # If no sentences were created, treat entire text as one sentence
        if not sentences and text.strip():
            sentences.append(Sentence(
                sentence_id=f"{para_id}_s_000",
                paragraph_id=para_id,
                sentence_index=0,
                text=text,
                start_char=0,
                end_char=len(text)
            ))

        return sentences

    def _extract_references(self, paragraphs: list[Paragraph]) -> list[Reference]:
        """Extract references from bibliography section."""
        references = []
        ref_index = 0
        in_references = False

        for para in paragraphs:
            # Check if we're entering references section
            if self.references_pattern.match(para.text):
                in_references = True
                continue

            # Extract references
            if in_references:
                # Skip empty paragraphs
                if not para.text.strip():
                    continue

                # Each non-empty paragraph in references section is a reference
                references.append(Reference(
                    reference_id=f"ref_{ref_index:03d}",
                    reference_index=ref_index,
                    raw_text=para.text
                ))
                ref_index += 1

        return references

    def _extract_metadata(self, doc: Document, paragraphs: list[Paragraph]) -> DocumentMetadata:
        """Extract document metadata."""
        # Word count and character count from paragraphs
        total_text = ' '.join(p.text for p in paragraphs)
        word_count = len(total_text.split())
        character_count = len(total_text)

        # Try to get author from document properties
        author = None
        created_date = None
        modified_date = None

        try:
            if hasattr(doc.core_properties, 'author'):
                author = doc.core_properties.author
            if hasattr(doc.core_properties, 'created'):
                created_date = doc.core_properties.created
            if hasattr(doc.core_properties, 'modified'):
                modified_date = doc.core_properties.modified
        except Exception as e:
            logger.warning("metadata_extraction_error", error=str(e))

        # Page count approximation (Word doesn't store this directly)
        page_count = max(1, word_count // 500)

        return DocumentMetadata(
            page_count=page_count,
            word_count=word_count,
            character_count=character_count,
            author=author,
            created_date=created_date,
            modified_date=modified_date
        )

    def _extract_title(self, doc: Document, sections: list[Section]) -> str:
        """Extract or infer document title."""
        # Try to get from document properties
        try:
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                return doc.core_properties.title
        except Exception:
            pass

        # Use first heading if available
        if sections and sections[0].section_title:
            return sections[0].section_title

        # Use first paragraph if it's short enough
        if doc.paragraphs and len(doc.paragraphs[0].text) < 100:
            return doc.paragraphs[0].text.strip()

        return "Untitled Document"

    def _get_xml_path(self, para: DocxParagraph) -> str | None:
        """Get the XML path to the paragraph element for export."""
        try:
            if hasattr(para, '_element'):
                element = para._element
                return f"//w:p[@id='{id(element)}']"
        except Exception as e:
            logger.warning("xml_path_extraction_error", error=str(e))

        return None


async def parse_docx(file_path: Path, title_override: str | None = None) -> DocObj:
    """Convenience function to parse a DOCX file."""
    parser = DocxParser()
    return await parser.parse(file_path, title_override)
